"""
Document Parser Service
Parse PDF, DOCX, TXT files into plain text for RAG indexing.
"""

import io
import os
import tempfile
from typing import Optional


class DocumentParser:
    """Parse various document formats into plain text."""

    SUPPORTED_MIMETYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/vnd.google-apps.document': 'gdoc',
        'text/plain': 'txt',
        'text/csv': 'txt',
        'text/markdown': 'txt',
    }

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.txt': 'txt',
        '.csv': 'txt',
        '.md': 'txt',
    }

    @classmethod
    def parse_file(cls, file_path: str, mime_type: Optional[str] = None) -> str:
        """
        Parse a file into plain text.
        
        Args:
            file_path: Path to the file
            mime_type: Optional MIME type hint
            
        Returns:
            Extracted plain text
        """
        # Determine format from mime_type or file extension
        fmt = None
        if mime_type and mime_type in cls.SUPPORTED_MIMETYPES:
            fmt = cls.SUPPORTED_MIMETYPES[mime_type]
        
        if not fmt:
            ext = os.path.splitext(file_path)[1].lower()
            fmt = cls.SUPPORTED_EXTENSIONS.get(ext)

        if not fmt:
            raise ValueError(f"Unsupported file format: {file_path} (mime: {mime_type})")

        if fmt == 'pdf':
            return cls._parse_pdf(file_path)
        elif fmt == 'docx':
            return cls._parse_docx(file_path)
        elif fmt == 'txt':
            return cls._parse_txt(file_path)
        else:
            raise ValueError(f"Unknown format: {fmt}")

    @classmethod
    def parse_bytes(cls, content: bytes, filename: str, mime_type: Optional[str] = None) -> str:
        """
        Parse file content (bytes) into plain text.
        
        Args:
            content: Raw file bytes
            filename: Original filename (for extension detection)
            mime_type: Optional MIME type hint
            
        Returns:
            Extracted plain text
        """
        # Write to temp file then parse
        ext = os.path.splitext(filename)[1] or '.tmp'
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            return cls.parse_file(tmp_path, mime_type)
        finally:
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Parse PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF is required for PDF parsing. Install: pip install PyMuPDF")

        text_parts = []
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text.strip()}")
            doc.close()
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF '{file_path}': {e}")

        return "\n\n".join(text_parts)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Parse DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install: pip install python-docx")

        text_parts = []
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX '{file_path}': {e}")

        return "\n\n".join(text_parts)

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """Parse plain text file."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise RuntimeError(f"Failed to read text file '{file_path}' with any encoding")

    @classmethod
    def is_supported(cls, filename: str, mime_type: Optional[str] = None) -> bool:
        """Check if a file format is supported for parsing."""
        if mime_type and mime_type in cls.SUPPORTED_MIMETYPES:
            return True
        ext = os.path.splitext(filename)[1].lower()
        return ext in cls.SUPPORTED_EXTENSIONS
